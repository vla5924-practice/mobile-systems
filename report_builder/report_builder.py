from argparse import ArgumentParser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from pathlib import Path
import os


parser = ArgumentParser()
parser.add_argument("--output", "-o", type=str, help="output document")
parser.add_argument("--template", "-t", type=str, help="template document")
parser.add_argument("--dir", "-d", type=str, help="path to qml files")
args = parser.parse_args()

tpl = open(args.template, "rb")
document = Document(tpl)
tpl.close()

doc = Document(args.template)
doc.add_heading("Report")

for path in Path(args.dir).glob("*.qml"):
    f = open(str(path), "r")
    doc.add_heading(os.path.basename(str(path)), level=2)
    para = doc.add_paragraph(f.read())
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    f.close()

doc.save(args.output)
